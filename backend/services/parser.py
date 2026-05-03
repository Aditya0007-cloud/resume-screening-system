from pathlib import Path

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

from backend.utils.text import extract_candidate_name, normalize_text


class ResumeParseError(ValueError):
    pass


def parse_resume(path: Path, original_filename: str) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        raw_text = _extract_pdf(path)
    elif suffix == ".docx":
        raw_text = _extract_docx(path)
    else:
        raise ResumeParseError("Only PDF and DOCX resumes are supported.")

    text = normalize_text(raw_text)
    if len(text) < 40:
        raise ResumeParseError(f"{original_filename} did not contain enough readable text.")
    return text, extract_candidate_name(raw_text, original_filename)


def _extract_pdf(path: Path) -> str:
    text_parts: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
    except Exception:
        text_parts = []

    text = "\n".join(text_parts).strip()
    if text:
        return text

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)
